import streamlit as st
import pandas as pd
import datetime
import io
import json
import os
import base64
from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
import pyotp
import qrcode

import time

st.set_page_config(page_title="Attendance Checker", page_icon="🕒", layout="wide")

is_auth = st.session_state.get('authenticated', False)
page_id = "dashboard" if is_auth else "gateway"

st.markdown("""
<style>
/* --- Robust Curtain Animation --- */
@keyframes curtainLeft {
    0% { right: 50%; opacity: 1; z-index: 9999999; }
    99% { right: 100%; opacity: 0; z-index: 9999999; }
    100% { right: 100%; opacity: 0; z-index: -1; visibility: hidden; }
}
@keyframes curtainRight {
    0% { left: 50%; opacity: 1; z-index: 9999999; }
    99% { left: 100%; opacity: 0; z-index: 9999999; }
    100% { left: 100%; opacity: 0; z-index: -1; visibility: hidden; }
}

/* Dynamic curtain classes defined below */

/* Smoothen the entrance of the main content */
.block-container {
    animation: slideUpFade 1.0s ease forwards;
    animation-delay: 0.1s;
    opacity: 0;
}
@keyframes slideUpFade {
    0% { opacity: 0; transform: translateY(30px); }
    100% { opacity: 1; transform: translateY(0); }
}

/* File Uploader Hover Glow */
[data-testid="stFileUploader"] {
    border: 1px solid rgba(0, 198, 255, 0.2);
    border-radius: 12px;
    background: rgba(0, 114, 255, 0.02);
    padding: 10px;
    transition: all 0.3s ease;
}
[data-testid="stFileUploader"]:hover {
    border-color: #00c6ff;
    background: rgba(0, 198, 255, 0.05);
    box-shadow: 0 4px 15px rgba(0, 198, 255, 0.1);
}

/* Premium Card Hover Effects */
div[data-testid="stExpander"] {
    background-color: rgba(255, 255, 255, 0.02);
    border-radius: 8px;
    border: 1px solid rgba(255, 255, 255, 0.1);
    transition: all 0.3s ease;
}
div[data-testid="stExpander"]:hover {
    border-color: rgba(255, 255, 255, 0.3);
    box-shadow: 0 4px 15px rgba(0, 0, 0, 0.5);
}

/* Premium Primary Button */
.stButton>button {
    background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
    color: white !important;
    border-radius: 8px;
    border: none;
    font-weight: bold;
    transition: transform 0.2s, box-shadow 0.2s;
}
.stButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 15px rgba(42, 82, 152, 0.4);
}

/* Premium DataFrame border */
div[data-testid="stDataFrame"] {
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 8px;
    overflow: hidden;
    box-shadow: 0 4px 20px rgba(0,0,0,0.2) !important;
}

/* --- Output Pop-up Animations --- */
@keyframes outputPopUp {
    0%, 30% { transform: scale(0.9) translateY(40px); opacity: 0; }
    100% { transform: scale(1) translateY(0); opacity: 1; }
}
div[data-testid="stDataFrame"], .output-alert-box {
    animation: outputPopUp 1.2s cubic-bezier(0.175, 0.885, 0.32, 1.275) forwards;
}

/* Make headers look slightly crisper */
h1, h2, h3 {
    font-family: "Inter", sans-serif;
    letter-spacing: -0.5px;
}
</style>
""" + f"""
<style>
.curtain-panel-left-{page_id} {{
    position: fixed; top: 0; left: 0; right: 50%; bottom: 0;
    background-color: #0e1117; border-right: 2px solid rgba(0, 198, 255, 0.4);
    box-shadow: 5px 0 20px rgba(0,0,0,0.5);
    animation: curtainLeft 1.2s cubic-bezier(0.86, 0, 0.07, 1) forwards;
    animation-delay: 0.1s; z-index: 9999999; pointer-events: none;
}}
.curtain-panel-right-{page_id} {{
    position: fixed; top: 0; left: 50%; right: 0; bottom: 0;
    background-color: #0e1117; border-left: 2px solid rgba(0, 198, 255, 0.4);
    box-shadow: -5px 0 20px rgba(0,0,0,0.5);
    animation: curtainRight 1.2s cubic-bezier(0.86, 0, 0.07, 1) forwards;
    animation-delay: 0.1s; z-index: 9999999; pointer-events: none;
}}
</style>
<div class="curtain-panel-left-{page_id}"></div>
<div class="curtain-panel-right-{page_id}"></div>
""", unsafe_allow_html=True)

def get_header_html(logo_path, title_text, emoji_fallback, width=40, motto=""):
    title_style = "background: linear-gradient(135deg, #00c6ff 0%, #0072ff 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; font-weight: 800; letter-spacing: -1px; line-height: 1.1;"
    
    motto_html = ""
    if motto:
        motto_html = f'<div style="margin-top: 8px; font-size: 1.15rem; background: linear-gradient(90deg, #89f7fe 0%, #66a6ff 100%); -webkit-background-clip: text; -webkit-text-fill-color: transparent; border-left: 2px solid #89f7fe; padding-left: 12px; font-weight: 600; letter-spacing: 0.5px; display: inline-block;">✦ {motto}</div>'

    if os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode()
        # Flexbox layout with optional motto cleanly stacked under the title
        img_html = f'<img src="data:image/jpeg;base64,{b64}" width="{width}" style="margin-right: 18px; border-radius: 8px; box-shadow: 0 4px 15px rgba(0, 114, 255, 0.4); flex-shrink: 0; margin-top: 5px;">'
        return f'<div style="display: flex; align-items: flex-start; margin-bottom: 2.5rem;">{img_html}<div style="display: flex; flex-direction: column; justify-content: center;"><h1 style="margin: 0; padding: 0; {title_style}">{title_text}</h1>{motto_html}</div></div>'
    return f'<div style="margin-bottom: 2.5rem;"><h1 style="{title_style} margin: 0 0 10px 0;">{emoji_fallback} {title_text}</h1>{motto_html}</div>'

# --- Security Gateway (Google Authenticator) ---
if 'totp_secret' not in st.session_state:
    st.session_state['totp_secret'] = 'JBSWY3DPEBLW64TN'  # Can be moved to st.secrets later

if not st.session_state.get('authenticated', False):
    st.markdown(get_header_html(
        "sentrium_logo.jpg", 
        "Sentrium Security Gateway", 
        "🛡️", 
        width=55, 
        motto="We protect and strengthen the digital foundation of businesses"
    ), unsafe_allow_html=True)
    
    st.info("Please enter your 6-digit verification code from Google Authenticator.")
    
    totp = pyotp.TOTP(st.session_state['totp_secret'])
    
    with st.form("totp_form"):
        user_code = st.text_input("Enter 6-digit code", type="password")
        submit_button = st.form_submit_button("Verify")
        
        if submit_button:
            if totp.verify(user_code):
                st.session_state['authenticated'] = True
                st.rerun()
            else:
                st.error("Invalid code. Please try again.")

    st.stop()

st.markdown(get_header_html("sentrium_logo.jpg", "Sentrium Attendance Tracker", "🕒", width=60), unsafe_allow_html=True)

st.markdown(
    """<div style="background-color: rgba(0, 114, 255, 0.05); border-left: 4px solid #0072ff; padding: 15px 20px; border-radius: 0 8px 8px 0; margin-bottom: 2rem; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
    <span style="color: #e2e8f0; font-size: 1.05rem; line-height: 1.5;">Upload your access control logs alongside the <strong style="color: #00c6ff;">Staff Roster</strong> and/or the <strong style="color: #00c6ff;">SOC Schedule</strong> to generate a full attendance report — covering late arrivals and absentees for both departments.</span>
    </div>""", unsafe_allow_html=True
)

# ── Three-column upload layout ─────────────────────────────────────────────
up_col1, up_col2, up_col3 = st.columns(3)

with up_col1:
    uploaded_files = st.file_uploader(
        "Access Control Logs",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=True,
        help="Upload the raw badge-in logs from the door access system"
    )

with up_col2:
    roster_file = st.file_uploader(
        "Standard Staff Master List",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=False,
        help="Upload the monthly roster for standard staff (Mon–Fri). Enables Absentee tracking."
    )

with up_col3:
    soc_file = st.file_uploader(
        "SOC Team Schedule",
        type=['csv', 'xlsx', 'xls'],
        accept_multiple_files=False,
        help="Upload the SOC team's rotating shift schedule. Enables SOC-specific Late & Absent reporting."
    )

st.markdown("<br>", unsafe_allow_html=True)

# ── SOC Shift time mapping ──────────────────────────────────────────────────
SOC_SHIFT_TIMES = {
    'morning':   datetime.time(7, 0, 0),
    'afternoon': datetime.time(14, 0, 0),
    'night':     datetime.time(19, 0, 0),
}

# ── Helper: compact export dropdown ────────────────────────────────────────
def render_export_dropdown(dataframe, period_label, report_title, report_subtitle, col_widths_pdf, key_prefix):
    """Renders a compact dropdown with Excel/Word/PDF export options."""
    with st.popover("Export Report"):
        # Excel
        buf_xlsx = io.BytesIO()
        with pd.ExcelWriter(buf_xlsx, engine='openpyxl') as writer:
            dataframe.to_excel(writer, index=False, sheet_name=report_title[:31])
        st.download_button(
            label="Download as Excel (.xlsx)",
            data=buf_xlsx.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"xlsx_{key_prefix}_{period_label}"
        )

        # Word
        doc = Document()
        doc.add_heading(f'{report_title} — {period_label}', level=1)
        doc.add_paragraph(report_subtitle)
        tbl = doc.add_table(rows=1, cols=len(dataframe.columns), style='Light Grid Accent 1')
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = tbl.rows[0].cells
        for i, cn in enumerate(dataframe.columns):
            hdr[i].text = str(cn)
        for _, row in dataframe.iterrows():
            rc = tbl.add_row().cells
            for i, cn in enumerate(dataframe.columns):
                rc[i].text = str(row[cn])
        buf_docx = io.BytesIO()
        doc.save(buf_docx)
        st.download_button(
            label="Download as Word (.docx)",
            data=buf_docx.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"docx_{key_prefix}_{period_label}"
        )

        # PDF
        pdf = FPDF()
        pdf.add_page(orientation='L')
        pdf.set_font('Helvetica', 'B', 14)
        pdf.cell(0, 10, f'{report_title} - {period_label}', new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.set_font('Helvetica', '', 9)
        pdf.cell(0, 8, report_subtitle, new_x="LMARGIN", new_y="NEXT", align='C')
        pdf.ln(5)
        pdf.set_font('Helvetica', 'B', 8)
        for i, h in enumerate(dataframe.columns):
            pdf.cell(col_widths_pdf[i] if i < len(col_widths_pdf) else 40, 8, str(h), border=1, align='C')
        pdf.ln()
        pdf.set_font('Helvetica', '', 7)
        for _, row in dataframe.iterrows():
            for i, cn in enumerate(dataframe.columns):
                w = col_widths_pdf[i] if i < len(col_widths_pdf) else 40
                pdf.cell(w, 7, str(row[cn]), border=1, align='C' if i in [0, 2] else 'L')
            pdf.ln()
        buf_pdf = io.BytesIO()
        pdf.output(buf_pdf)
        st.download_button(
            label="Download as PDF (.pdf)",
            data=buf_pdf.getvalue(),
            file_name=f"{key_prefix}_{period_label.replace(' ', '_')}.pdf",
            mime="application/pdf",
            key=f"pdf_{key_prefix}_{period_label}"
        )


# ── Helper: period filter (All / Daily / Weekly / Monthly) ──────────────────
def period_filter(df, date_col, parsed_time_col, radio_key):
    period = st.radio("View period:", ["All", "Daily", "Weekly", "Monthly"], horizontal=True, key=radio_key)
    if period == "Daily":
        df['_date_label'] = df[parsed_time_col].dt.date
        opts = sorted(df['_date_label'].unique(), reverse=True)
        sel = st.selectbox("Select Date", opts, key=f"{radio_key}_sel")
        return df[df['_date_label'] == sel].copy(), f"Day {sel}"
    elif period == "Weekly":
        df['_week'] = df[parsed_time_col].dt.isocalendar().year.astype(str) + '-W' + \
                      df[parsed_time_col].dt.isocalendar().week.astype(str).str.zfill(2)
        opts = sorted(df['_week'].unique(), reverse=True)
        sel = st.selectbox("Select Week", opts, key=f"{radio_key}_sel")
        return df[df['_week'] == sel].copy(), f"Week {sel}"
    elif period == "Monthly":
        df['_month'] = df[parsed_time_col].dt.to_period('M').astype(str)
        opts = sorted(df['_month'].unique(), reverse=True)
        sel = st.selectbox("Select Month", opts, key=f"{radio_key}_sel")
        return df[df['_month'] == sel].copy(), f"Month {sel}"
    return df.copy(), "All Time"


if uploaded_files:
    with st.spinner("Decrypting Access Logs & Compiling Report..."):
        time.sleep(1.2)

    try:
        # ── STEP 1: Load & normalise access logs ────────────────────────────
        STANDARD_COLS = {
            'time': 'Time', 'door name': 'Door Name',
            'event description': 'Event Description', 'personnel id': 'Personnel ID',
            'first name': 'First Name', 'last name': 'Last Name', 'door number': 'Door Number',
        }

        def normalize_columns(df_in):
            return df_in.rename(columns={
                c: STANDARD_COLS.get(str(c).strip().lower(), str(c).strip())
                for c in df_in.columns
            })

        def has_required_cols(cols):
            lc = [str(c).strip().lower() for c in cols]
            return all(t in lc for t in ['time', 'door name', 'event description'])

        all_dfs = []
        for uf in uploaded_files:
            if uf.name.endswith('.csv'):
                tmp = pd.read_csv(uf)
                tmp.columns = tmp.columns.astype(str).str.strip()
                all_dfs.append(normalize_columns(tmp))
            else:
                sheets = pd.read_excel(uf, sheet_name=None)
                for sname, d in sheets.items():
                    if d.empty:
                        continue
                    d.columns = d.columns.astype(str).str.strip()
                    if not has_required_cols(d.columns):
                        found = False
                        for i in range(min(20, len(d))):
                            rv = [str(v).strip() for v in d.iloc[i].values]
                            if has_required_cols(rv):
                                d.columns = rv
                                d = d.iloc[i+1:].reset_index(drop=True)
                                found = True
                                break
                        if not found:
                            continue
                    d = d.dropna(how='all').reset_index(drop=True)
                    d = normalize_columns(d)
                    if 'Time' in d.columns and 'Door Name' in d.columns:
                        all_dfs.append(d)

        if not all_dfs:
            st.error("No valid data found in the uploaded log files.")
            st.stop()

        df = pd.concat(all_dfs, ignore_index=True)

        # Define helper to clean Personnel ID and prevent merge float/str errors
        def clean_id(val):
            if pd.isna(val):
                return ""
            val_str = str(val).strip()
            if val_str.endswith('.0'):
                val_str = val_str[:-2]
            return val_str

        # Validate
        for col in ['Event Description', 'Time', 'Door Name']:
            if col not in df.columns:
                st.error(f"Missing required column: **{col}**")
                st.stop()

        if 'Personnel ID' in df.columns:
            df['Personnel ID'] = df['Personnel ID'].apply(clean_id)

        # Parse timestamps
        df['Parsed_Time'] = pd.to_datetime(df['Time'], errors='coerce', dayfirst=True)
        nat_mask = df['Parsed_Time'].isna()
        if nat_mask.any():
            df.loc[nat_mask, 'Parsed_Time'] = pd.to_datetime(df.loc[nat_mask, 'Time'], errors='coerce', dayfirst=False)
        df = df.dropna(subset=['Parsed_Time']).copy()

        df['Event Description'] = df['Event Description'].astype(str).str.strip()
        df['Door Name'] = df['Door Name'].astype(str).str.strip()

        # Fixed filters
        TARGET_DOORS  = ['Main entrance Ground Flr', 'Main entrance Upfloor']
        TARGET_EVENTS = ['Password', 'Normal Open']
        df_entry = df[
            df['Event Description'].str.lower().isin([e.lower() for e in TARGET_EVENTS]) &
            df['Door Name'].str.lower().isin([d.lower() for d in TARGET_DOORS])
        ].copy()

        df_entry['Date']      = df_entry['Parsed_Time'].dt.date
        df_entry['Time_Only'] = df_entry['Parsed_Time'].dt.time
        df_entry['Date_dt']   = pd.to_datetime(df_entry['Date'])

        # First badge-in per person per day
        df_entry = df_entry.sort_values('Parsed_Time')
        df_first = df_entry.drop_duplicates(subset=['Date', 'Personnel ID'], keep='first').copy()

        # ── STEP 2: Load rosters ────────────────────────────────────────────
        def load_simple_excel(file_obj):
            if file_obj.name.endswith('.csv'):
                return pd.read_csv(file_obj)
            return pd.read_excel(file_obj)

        df_roster = None
        soc_ids_from_roster = set()
        if roster_file:
            df_roster = load_simple_excel(roster_file)
            df_roster.columns = df_roster.columns.astype(str).str.strip()
            # Expect exactly: Personnel ID, Expected Days
            # Normalize known column names
            col_map = {}
            for c in df_roster.columns:
                cl = c.lower()
                if 'personnel' in cl or cl == 'id':
                    col_map[c] = 'Personnel ID'
                elif 'expected' in cl or 'days' in cl:
                    col_map[c] = 'Expected Days'
            df_roster = df_roster.rename(columns=col_map)

            if 'Personnel ID' not in df_roster.columns or 'Expected Days' not in df_roster.columns:
                st.warning("Master List must contain 'Personnel ID' and 'Expected Days' columns. Absentee tracking disabled.")
                df_roster = None
            else:
                df_roster['Personnel ID'] = df_roster['Personnel ID'].apply(clean_id)
                df_roster['Expected Days'] = df_roster['Expected Days'].astype(str).str.strip()
                # Identify SOC staff (Expected Days = "SOC")
                soc_mask = df_roster['Expected Days'].str.upper() == 'SOC'
                soc_ids_from_roster = set(df_roster.loc[soc_mask, 'Personnel ID'].unique())
                # Keep only standard staff for absentee logic
                df_std_roster = df_roster[~soc_mask].copy()
                df_std_roster['Days_Per_Week'] = pd.to_numeric(df_std_roster['Expected Days'], errors='coerce')
                df_std_roster = df_std_roster.dropna(subset=['Days_Per_Week']).copy()
                df_std_roster['Days_Per_Week'] = df_std_roster['Days_Per_Week'].astype(int)

        df_soc = None
        if soc_file:
            df_soc = load_simple_excel(soc_file)
            df_soc.columns = df_soc.columns.astype(str).str.strip()
            df_soc = normalize_columns(df_soc)
            # Expect: Personnel ID, Date, Shift (+ optional First Name, Last Name)
            for needed in ['Personnel ID', 'Date', 'Shift']:
                if needed not in df_soc.columns:
                    st.warning(f"SOC Schedule must contain a '{needed}' column. SOC tab disabled.")
                    df_soc = None
                    break
            if df_soc is not None:
                df_soc['Date'] = pd.to_datetime(df_soc['Date'], errors='coerce', dayfirst=True).dt.date
                df_soc = df_soc.dropna(subset=['Date']).copy()
                df_soc['Shift_Lower'] = df_soc['Shift'].astype(str).str.strip().str.lower()
                df_soc['Shift_Start'] = df_soc['Shift_Lower'].map(SOC_SHIFT_TIMES)
                df_soc['Personnel ID'] = df_soc['Personnel ID'].apply(clean_id)
                fn_cols = [c for c in ['First Name', 'Last Name'] if c in df_soc.columns]
                df_soc['Full Name'] = df_soc[fn_cols].fillna('').astype(str).apply(
                    lambda r: ' '.join(r).strip(), axis=1) if fn_cols else ''

        # ── STEP 3: Tabs ────────────────────────────────────────────────────
        st.divider()
        tab_std, tab_soc = st.tabs(["Standard Staff", "SOC Team"])


        # ════════════════════════════════════════════════════════════════════
        # TAB 1 — STANDARD STAFF
        # ════════════════════════════════════════════════════════════════════
        with tab_std:

            CUTOFF = datetime.time(8, 30, 0)
            WORKDAYS = {0, 1, 2, 3, 4}  # Mon–Fri

            # Combine SOC IDs from master list + SOC schedule
            soc_ids = soc_ids_from_roster.copy()
            if df_soc is not None:
                soc_ids.update(df_soc['Personnel ID'].astype(str).str.strip().unique())

            df_std_first = df_first[~df_first['Personnel ID'].astype(str).str.strip().isin(soc_ids)].copy()

            # ── Late Arrivals ────────────────────────────────────────────────
            st.subheader("Late Arrivals — Standard Staff")

            # Filter Mon-Fri only
            df_std_workday = df_std_first[df_std_first['Date_dt'].dt.dayofweek.isin(WORKDAYS)].copy()
            df_late = df_std_workday[df_std_workday['Time_Only'] >= CUTOFF].copy()

            # Build name columns
            name_cols = [c for c in ['First Name', 'Last Name'] if c in df_late.columns]
            df_late['Full Name'] = df_late[name_cols].fillna('').astype(str).apply(
                lambda r: ' '.join(r).strip(), axis=1) if name_cols else ''
            df_late['Date_Str'] = df_late['Date'].apply(lambda d: d.strftime('%Y-%m-%d'))
            df_late['Time_Str'] = df_late['Time_Only'].apply(lambda t: t.strftime('%H:%M:%S'))

            if df_late.empty:
                st.markdown(
                    '<div class="output-alert-box" style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                    '<h3 style="margin:0;color:#30d158;">Perfect Attendance</h3>'
                    '<p style="margin:5px 0 0;color:#e2e8f0;">No standard staff checked in after 8:30 AM.</p></div>',
                    unsafe_allow_html=True
                )
            else:
                st.markdown(
                    f'<div class="output-alert-box" style="background:rgba(255,69,58,0.1);border-left:5px solid #ff453a;padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                    f'<h3 style="margin:0;color:#ff6b6b;">{len(df_late)} Late Check-in Record(s)</h3>'
                    f'<p style="margin:5px 0 0;color:#e2e8f0;">Standard staff who arrived at or after <strong>8:30 AM</strong> on a working day.</p></div>',
                    unsafe_allow_html=True
                )

                # Period filter row + export dropdown on the right
                pf_col, ex_col = st.columns([4, 1])
                with pf_col:
                    df_late_f, period_label_late = period_filter(df_late.copy(), 'Date', 'Parsed_Time', 'std_late_period')

                if df_late_f.empty:
                    st.info("No late arrivals for the selected period.")
                else:
                    # Aggregate: who was late, on which dates, at what times
                    late_agg = (
                        df_late_f.sort_values(['Date', 'Time_Only'])
                        .groupby(['Personnel ID', 'Full Name'])
                        .agg(
                            Times_Late=('Date', 'nunique'),
                            Dates_Late=('Date_Str', lambda x: ', '.join(x)),
                            Times_In=('Time_Str', lambda x: ', '.join(x)),
                        )
                        .reset_index()
                        .sort_values('Times_Late', ascending=False)
                    )

                    # Attach Expected Days/Week from roster if available
                    if df_roster is not None:
                        late_agg = late_agg.merge(
                            df_std_roster[['Personnel ID', 'Days_Per_Week']],
                            on='Personnel ID', how='left'
                        )
                        late_agg['Days_Per_Week'] = late_agg['Days_Per_Week'].fillna('–')
                        late_agg = late_agg[['Personnel ID', 'Full Name', 'Days_Per_Week',
                                            'Times_Late', 'Dates_Late', 'Times_In']]
                        late_agg.columns = ['Personnel ID', 'Full Name', 'Expected/Week',
                                            'Times Late', 'Dates Late', 'Times In']
                    else:
                        late_agg = late_agg[['Personnel ID', 'Full Name', 'Times_Late', 'Dates_Late', 'Times_In']]
                        late_agg.columns = ['Personnel ID', 'Full Name', 'Times Late', 'Dates Late', 'Times In']

                    st.dataframe(late_agg, hide_index=True, use_container_width=True)

                    with ex_col:
                        render_export_dropdown(
                            late_agg, period_label_late,
                            "Standard Staff Late Arrivals",
                            "Personnel who checked in at or after 08:30 AM on working days.",
                            [20, 40, 15, 15, 80, 80], "std_late"
                        )

            # ── Weekly Attendance Compliance ──────────────────────────────────
            st.divider()
            st.subheader("Weekly Attendance Compliance — Standard Staff")

            if df_roster is None:
                st.info("Upload the Standard Staff Master List to enable weekly compliance tracking.")
            else:
                # Build name lookup {Personnel ID -> Full Name}
                _name_lkp = {}
                if 'First Name' in df_first.columns:
                    _nl = df_first[['Personnel ID', 'First Name', 'Last Name']].copy()
                    _nl['Personnel ID'] = _nl['Personnel ID'].astype(str).str.strip()
                    _nc_cols = [c for c in ['First Name', 'Last Name'] if c in _nl.columns]
                    _nl['Full Name'] = _nl[_nc_cols].fillna('').astype(str).apply(
                        lambda r: ' '.join(r).strip(), axis=1)
                    _nl = _nl.drop_duplicates(subset='Personnel ID')
                    _name_lkp = dict(zip(_nl['Personnel ID'], _nl['Full Name']))

                # Workday-only (Mon–Fri) badge-ins for standard staff
                _df_att = df_std_first[df_std_first['Date_dt'].dt.dayofweek.isin(WORKDAYS)].copy()
                _df_att['Personnel ID'] = _df_att['Personnel ID'].astype(str).str.strip()
                _df_att['ISO_Year_n'] = _df_att['Parsed_Time'].dt.isocalendar().year.astype(int)
                _df_att['ISO_Week_n'] = _df_att['Parsed_Time'].dt.isocalendar().week.astype(int)
                _df_att['Week_Label'] = (_df_att['ISO_Year_n'].astype(str) + '-W' +
                                         _df_att['ISO_Week_n'].astype(str).str.zfill(2))
                _df_att['Date_Str'] = _df_att['Date'].apply(lambda d: d.strftime('%Y-%m-%d'))
                _df_att['Time_Str'] = _df_att['Time_Only'].apply(lambda t: t.strftime('%H:%M'))

                if _df_att.empty:
                    st.info("No workday attendance data found in the uploaded logs.")
                else:
                    _min_dt = _df_att['Parsed_Time'].min()
                    _max_dt = _df_att['Parsed_Time'].max()

                    # All ISO week Mondays in the log date range
                    _wk_starts = pd.date_range(
                        _min_dt.to_period('W').to_timestamp(),
                        _max_dt.to_period('W').to_timestamp(),
                        freq='W-MON'
                    )
                    _iso_weeks = []
                    for _ws in _wk_starts:
                        _ic = _ws.isocalendar()
                        _iso_weeks.append({
                            'year': int(_ic[0]), 'week': int(_ic[1]),
                            'label': f"{int(_ic[0])}-W{str(int(_ic[1])).zfill(2)}"
                        })
                    if not _iso_weeks:
                        _ic = _min_dt.isocalendar()
                        _iso_weeks = [{'year': int(_ic[0]), 'week': int(_ic[1]),
                                       'label': f"{int(_ic[0])}-W{str(int(_ic[1])).zfill(2)}"}]

                    # Build per-person per-week compliance rows
                    _comp_rows = []
                    for _, _pr in df_std_roster.iterrows():
                        _pid  = str(_pr['Personnel ID'])
                        _exp  = int(_pr['Days_Per_Week'])
                        _fname = _name_lkp.get(_pid, '')
                        _p_att = _df_att[_df_att['Personnel ID'] == _pid]

                        for _wk in _iso_weeks:
                            try:
                                _mon = datetime.date.fromisocalendar(_wk['year'], _wk['week'], 1)
                            except Exception:
                                continue
                            # All 5 Mon–Fri calendar dates in this ISO week
                            _wd_strs = [
                                (_mon + datetime.timedelta(days=_d)).strftime('%Y-%m-%d')
                                for _d in range(5)
                            ]
                            _wk_att = _p_att[_p_att['Week_Label'] == _wk['label']].sort_values('Date')
                            _att_strs = set(_wk_att['Date_Str'].tolist())
                            # "YYYY-MM-DD (HH:MM)" for each day attended
                            _att_entries = [
                                f"{_r['Date_Str']} ({_r['Time_Str']})"
                                for _, _r in _wk_att.iterrows()
                            ]
                            # Workdays with NO badge-in recorded
                            _absent = [_d for _d in _wd_strs if _d not in _att_strs]
                            _actual  = len(_att_strs)
                            _deficit = max(0, _exp - _actual)
                            _comp_rows.append({
                                'Personnel ID':         _pid,
                                'Full Name':            _fname,
                                'Week':                 _wk['label'],
                                'Expected Days':        _exp,
                                'Days Present':         _actual,
                                'Deficit':              _deficit,
                                'Status':               '✅ Met' if _deficit == 0 else f'⚠️ Short by {_deficit}',
                                'Days Attended & Times': ', '.join(_att_entries) if _att_entries else '—',
                                'Absent Dates':         ', '.join(_absent) if _absent else '—',
                            })

                    _df_comp = pd.DataFrame(_comp_rows)

                    # Summary banner
                    _nc_count = _df_comp[_df_comp['Deficit'] > 0]['Personnel ID'].nunique()
                    if _nc_count == 0:
                        st.markdown(
                            '<div style="background:rgba(48,209,88,0.1);border-left:5px solid #30d158;'
                            'padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                            '<h3 style="margin:0;color:#30d158;">✅ Full Compliance</h3>'
                            '<p style="margin:5px 0 0;color:#e2e8f0;">All standard staff met their '
                            'required weekly attendance.</p></div>',
                            unsafe_allow_html=True
                        )
                    else:
                        st.markdown(
                            f'<div style="background:rgba(255,149,0,0.1);border-left:5px solid #ff9500;'
                            f'padding:15px 20px;border-radius:6px;margin-bottom:1.5rem;">'
                            f'<h3 style="margin:0;color:#ff9500;">⚠️ {_nc_count} staff member(s) had '
                            f'non-compliant weeks</h3>'
                            f'<p style="margin:5px 0 0;color:#e2e8f0;">Based on required days per week '
                            f'vs. actual badge-in records.</p></div>',
                            unsafe_allow_html=True
                        )

                    # Period selector + export
                    _pf2_col, _ex2_col = st.columns([4, 1])
                    _all_wk_opts = sorted(_df_comp['Week'].unique(), reverse=True)
                    with _pf2_col:
                        _comp_view = st.radio(
                            "View period:", ["All Weeks", "Select Week"],
                            horizontal=True, key="comp_period_radio"
                        )
                    if _comp_view == "Select Week" and _all_wk_opts:
                        _sel_wk = st.selectbox("Select Week", _all_wk_opts, key="comp_week_sel")
                        _detail_view = _df_comp[_df_comp['Week'] == _sel_wk].copy()
                    else:
                        _detail_view = _df_comp.copy()

                    st.markdown(
                        "**Weekly Detail** — badge-in dates & times attended, and Mon–Fri dates "
                        "absent, per person per week."
                    )
                    st.dataframe(_detail_view, hide_index=True, use_container_width=True)

                    # Per-person summary
                    st.markdown("**Summary** — aggregate compliance per employee.")
                    _nc_wk_list = (
                        _df_comp[_df_comp['Deficit'] > 0]
                        .groupby('Personnel ID')['Week']
                        .apply(lambda x: ', '.join(sorted(x)))
                        .reset_index()
                        .rename(columns={'Week': 'Non-Compliant Weeks'})
                    )
                    _summary = (
                        _df_comp.groupby(['Personnel ID', 'Full Name'])
                        .agg(
                            Expected_Per_Week=('Expected Days', 'first'),
                            Weeks_NC=('Deficit', lambda x: int((x > 0).sum())),
                            Total_Deficit=('Deficit', 'sum'),
                        )
                        .reset_index()
                    )
                    _summary = _summary.merge(_nc_wk_list, on='Personnel ID', how='left')
                    _summary['Non-Compliant Weeks'] = _summary['Non-Compliant Weeks'].fillna('—')
                    _summary = _summary.sort_values('Total_Deficit', ascending=False)
                    _summary.columns = [
                        'Personnel ID', 'Full Name', 'Expected/Week',
                        'Weeks Non-Compliant', 'Total Deficit Days', 'Non-Compliant Weeks'
                    ]
                    st.dataframe(_summary, hide_index=True, use_container_width=True)

                    with _ex2_col:
                        render_export_dropdown(
                            _detail_view, "All Time",
                            "Weekly Attendance Compliance",
                            "Standard staff weekly attendance vs. required days per week.",
                            [15, 35, 12, 12, 12, 8, 15, 60, 60], "std_compliance"
                        )


        # ════════════════════════════════════════════════════════════════════
        # TAB 2 — SOC TEAM
        # ════════════════════════════════════════════════════════════════════
        with tab_soc:
            if df_soc is None:
                st.info("Upload the SOC Team Schedule file to enable the SOC Department report.")
                st.markdown("""
                **Expected columns in the SOC Schedule file:**

                | Personnel ID | First Name | Last Name | Date | Shift |
                |---|---|---|---|---|
                | 201 | Jane | Smith | 14/04/2026 | Morning |
                | 201 | Jane | Smith | 15/04/2026 | Night |
                | 202 | James | Okoro | 14/04/2026 | Afternoon |

                **Shift options:** `Morning` (07:00) · `Afternoon` (14:00) · `Night` (19:00)
                > Days off should simply be omitted — do not add a row for days off.
                """)
            else:
                # Cross-reference SOC schedule with access logs
                _soc_pids = df_soc['Personnel ID'].unique()
                _soc_logs = df_first[df_first['Personnel ID'].astype(str).str.strip().isin(_soc_pids)].copy()
                _soc_logs['Personnel ID'] = _soc_logs['Personnel ID'].astype(str).str.strip()

                # Merge schedule → actual logs (left join keeps every scheduled shift)
                _merged = df_soc.merge(
                    _soc_logs[['Personnel ID', 'Date', 'Time_Only', 'Parsed_Time']],
                    on=['Personnel ID', 'Date'],
                    how='left'
                )

                # ── Build unified shift compliance table ────────────────────────
                _sc = _merged.copy()
                _sc['Date_Str']        = _sc['Date'].apply(
                    lambda d: d.strftime('%Y-%m-%d') if hasattr(d, 'strftime') else str(d))
                _sc['Sched_Start_Str'] = _sc['Shift_Start'].apply(
                    lambda t: t.strftime('%H:%M') if pd.notna(t) else '—')
                _sc['Actual_In_Str']   = _sc['Time_Only'].apply(
                    lambda t: t.strftime('%H:%M') if pd.notna(t) else '—')
                _sc['Shift_Title']     = _sc['Shift'].astype(str).str.strip().str.title()
                _sc['Date_dt_col']     = pd.to_datetime(_sc['Date_Str'])

                def _soc_status(row):
                    if pd.isna(row['Time_Only']):
                        return '🚫 Absent'
                    if pd.notna(row['Shift_Start']) and row['Time_Only'] > row['Shift_Start']:
                        return '⚠️ Late'
                    return '✅ On Time'

                def _soc_mins(row):
                    if pd.isna(row['Time_Only']) or pd.isna(row['Shift_Start']):
                        return 0
                    if row['Time_Only'] > row['Shift_Start']:
                        return int(
                            (datetime.datetime.combine(row['Date'], row['Time_Only']) -
                             datetime.datetime.combine(row['Date'], row['Shift_Start'])
                             ).total_seconds() / 60
                        )
                    return 0

                _sc['Status']    = _sc.apply(_soc_status, axis=1)
                _sc['Mins Late'] = _sc.apply(_soc_mins, axis=1)

                # Full compliance table (one row per scheduled shift per person)
                _soc_comp_all = (
                    _sc[['Personnel ID', 'Full Name', 'Date_Str', 'Shift_Title',
                          'Sched_Start_Str', 'Actual_In_Str', 'Status', 'Mins Late']]
                    .rename(columns={
                        'Date_Str': 'Shift Date', 'Shift_Title': 'Shift',
                        'Sched_Start_Str': 'Scheduled Start',
                        'Actual_In_Str':   'Actual Check-in'
                    })
                    .sort_values(['Personnel ID', 'Shift Date'])
                )

                # ── Summary banner ──────────────────────────────────────────────
                _late_pids = _sc[_sc['Status'] == '⚠️ Late']['Personnel ID'].nunique()
                _abs_pids  = _sc[_sc['Status'] == '🚫 Absent']['Personnel ID'].nunique()
                _banner    = []
                if _late_pids:
                    _banner.append(f"<strong style='color:#ff9500;'>{_late_pids}</strong> member(s) arrived late")
                if _abs_pids:
                    _banner.append(f"<strong style='color:#ff453a;'>{_abs_pids}</strong> member(s) had absences")

                if not _banner:
                    st.markdown(
                        '<div class="output-alert-box" style="background:rgba(48,209,88,0.1);'
                        'border-left:5px solid #30d158;padding:15px 20px;border-radius:6px;'
                        'margin-bottom:1.5rem;">'
                        '<h3 style="margin:0;color:#30d158;">✅ Perfect SOC Compliance!</h3>'
                        '<p style="margin:5px 0 0;color:#e2e8f0;">All SOC members were on time '
                        'for every scheduled shift.</p></div>',
                        unsafe_allow_html=True
                    )
                else:
                    st.markdown(
                        f'<div class="output-alert-box" style="background:rgba(255,69,58,0.1);'
                        f'border-left:5px solid #ff453a;padding:15px 20px;border-radius:6px;'
                        f'margin-bottom:1.5rem;">'
                        f'<h3 style="margin:0;color:#ff6b6b;">SOC Shift Compliance Issues</h3>'
                        f'<p style="margin:5px 0 0;color:#e2e8f0;">{" · ".join(_banner)}</p></div>',
                        unsafe_allow_html=True
                    )

                # ── Shift Compliance Detail ─────────────────────────────────────
                st.subheader("Shift Compliance — SOC Team")
                st.markdown(
                    "Every row is one **scheduled shift**. &nbsp;"
                    "✅ On Time &nbsp;·&nbsp; ⚠️ Late (arrived after shift start) "
                    "&nbsp;·&nbsp; 🚫 Absent (no badge-in recorded)."
                )

                _sc_pf_col, _sc_ex_col = st.columns([4, 1])
                with _sc_pf_col:
                    _sc_f, _sc_period = period_filter(
                        _sc.copy(), 'Date', 'Date_dt_col', 'soc_comp_period'
                    )

                if _sc_f.empty:
                    st.info("No SOC shifts found for the selected period.")
                else:
                    _sc_display = (
                        _sc_f[['Personnel ID', 'Full Name', 'Date_Str', 'Shift_Title',
                                'Sched_Start_Str', 'Actual_In_Str', 'Status', 'Mins Late']]
                        .rename(columns={
                            'Date_Str': 'Shift Date', 'Shift_Title': 'Shift',
                            'Sched_Start_Str': 'Scheduled Start',
                            'Actual_In_Str':   'Actual Check-in'
                        })
                        .sort_values(['Personnel ID', 'Shift Date'])
                    )
                    st.dataframe(_sc_display, hide_index=True, use_container_width=True)

                    with _sc_ex_col:
                        render_export_dropdown(
                            _sc_display, _sc_period,
                            "SOC Team Shift Compliance",
                            "SOC personnel shift attendance: On Time, Late, or Absent.",
                            [18, 40, 14, 14, 16, 16, 14, 12], "soc_compliance"
                        )

                # ── Per-Person Summary ──────────────────────────────────────────
                st.divider()
                st.subheader("Compliance Summary — SOC Team")
                st.markdown(
                    "Aggregated per person: total scheduled shifts, on-time count, "
                    "late arrivals, absences, and total minutes late."
                )

                _soc_summary = (
                    _soc_comp_all
                    .groupby(['Personnel ID', 'Full Name'])
                    .agg(
                        Scheduled_Shifts=('Shift Date', 'count'),
                        On_Time=('Status',    lambda x: (x == '✅ On Time').sum()),
                        Late=   ('Status',    lambda x: x.str.startswith('⚠️').sum()),
                        Absent= ('Status',    lambda x: (x == '🚫 Absent').sum()),
                        Total_Mins_Late=('Mins Late', 'sum'),
                    )
                    .reset_index()
                    .sort_values(['Absent', 'Late'], ascending=False)
                )
                _soc_summary.columns = [
                    'Personnel ID', 'Full Name', 'Scheduled Shifts',
                    'On Time', 'Late', 'Absent', 'Total Mins Late'
                ]
                st.dataframe(_soc_summary, hide_index=True, use_container_width=True)

                render_export_dropdown(
                    _soc_summary, "All Time",
                    "SOC Team Compliance Summary",
                    "Per-person SOC shift compliance overview.",
                    [18, 40, 20, 14, 14, 14, 20], "soc_summary"
                )

    except Exception as e:
        st.error(f"An error occurred while processing: {e}")
